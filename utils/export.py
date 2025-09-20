import io
import requests
from typing import Dict, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT_NAME = "Source Sans 3"


def _set_document_defaults(doc: Document):
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.font.size = Pt(11)


def _add_heading(doc: Document, text: str):
    h = doc.add_paragraph()
    run = h.add_run(text)
    run.bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_bullets(doc: Document, lines: List[str]):
    for line in lines:
        if not line.strip():
            continue
        p = doc.add_paragraph(style=None)
        p.style = doc.styles['List Paragraph']
        p.add_run(line.strip("- "))


def jd_to_docx(jd_sections: Dict[str, List[str]], title: str, tenant_name: str = "", logo_url: str = "", primary_colour: str = "#111827") -> bytes:
    doc = Document()
    _set_document_defaults(doc)

    # Header with optional logo and tenant
    if logo_url:
        try:
            img = requests.get(logo_url, timeout=5).content
            doc.add_picture(io.BytesIO(img), width=Inches(1.4))
        except Exception:
            pass

    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    if tenant_name:
        sub = doc.add_paragraph(tenant_name)
        sub.runs[0].font.size = Pt(11)

    # Sections in standard order when present
    for sec, body in jd_sections.items():
        doc.add_paragraph("")
        _add_heading(doc, sec)
        # Heuristic: list vs paragraph
        if len(body) > 1:
            _add_bullets(doc, [ln for ln in body if ln.strip()])
        else:
            p = doc.add_paragraph("\n".join(body).strip())
            p.paragraph_format.space_after = Pt(6)

    # Export bytes
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.getvalue()


def jd_to_markdown(jd_sections: Dict[str, List[str]], title: str) -> str:
    lines = [f"# {title}"]
    for sec, body in jd_sections.items():
        lines.append(f"\n## {sec}")
        if len(body) > 1:
            for b in body:
                if b.strip():
                    lines.append(f"- {b.strip('- ').strip()}")
        else:
            lines.append("\n" + "\n".join(body))
    return "\n".join(lines)